from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("reports/Huong_dan_cai_thien_du_an_chi_tiet.docx")
NAVY = "1F4E79"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
GRAY = "667085"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def font(run, size=10.5, bold=False, italic=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    properties.append(element)


def margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    cell_margins = properties.first_child_found_in("w:tcMar")
    if cell_margins is None:
        cell_margins = OxmlElement("w:tcMar")
        properties.append(cell_margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = cell_margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            cell_margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_text(cell, text, bold=False, color=None, size=9.2):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    font(run, size=size, bold=bold, color=color)
    margins(cell)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        if widths:
            cell.width = Inches(widths[index])
        shade(cell, NAVY)
        cell_text(cell, header, bold=True, color=WHITE)
    for row_number, values in enumerate(rows):
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            if widths:
                cell.width = Inches(widths[index])
            cell_text(cell, value)
            if row_number % 2:
                shade(cell, LIGHT_GRAY)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(document, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    font(paragraph.add_run(text), size=10.5)
    return paragraph


def add_number(document, title, detail):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    font(paragraph.add_run(f"{title}: "), size=10.5, bold=True, color=NAVY)
    font(paragraph.add_run(detail), size=10.5)


def callout(document, label, text, fill=PALE_BLUE):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), fill)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    font(paragraph.add_run(f"{label}: "), size=10, bold=True, color=NAVY)
    font(paragraph.add_run(text), size=10)
    margins(table.cell(0, 0), top=120, bottom=120, start=150, end=150)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def code_block(document, code):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), "F7F7F8")
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    font(paragraph.add_run(code.strip()), size=8.4, name="Consolas", color="202124")
    margins(table.cell(0, 0), top=120, bottom=120, start=150, end=150)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def checklist(document, items):
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(3)
        font(paragraph.add_run("☐ "), size=11, color=BLUE)
        font(paragraph.add_run(item), size=10.2)


document = Document()
section = document.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.35)

normal = document.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for style_name, size, before, after in (
    ("Heading 1", 16, 16, 8),
    ("Heading 2", 13, 12, 6),
    ("Heading 3", 11.5, 8, 4),
):
    style = document.styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE if style_name != "Heading 3" else NAVY)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("HCMC REAL ESTATE PRICE INTELLIGENCE  |  IMPROVEMENT GUIDE"), size=8, color=GRAY)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Hướng dẫn cải thiện dự án • 01/09/2026 • Trang "), size=8.5, color=GRAY)
page_field = OxmlElement("w:fldSimple")
page_field.set(qn("w:instr"), "PAGE")
footer._p.append(page_field)

# Memo masthead
spacer = document.add_paragraph()
spacer.paragraph_format.space_after = Pt(8)
title = document.add_paragraph()
title.paragraph_format.space_after = Pt(3)
font(title.add_run("HƯỚNG DẪN CẢI THIỆN DỰ ÁN"), size=24, bold=True, color=NAVY)
subtitle = document.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
font(subtitle.add_run("Từ portfolio ML đến hệ thống có tính đúng đắn, khả năng vận hành và giá trị kinh doanh"), size=13, italic=True, color=GRAY)
for label, value in (
    ("Dự án", "HCMC Real Estate Price Intelligence"),
    ("Phạm vi", "Problem → AI/ML correctness → Software Engineering → Production / Business value"),
    ("Mục tiêu", "Cải thiện theo bằng chứng, không bổ sung công nghệ chỉ để trình diễn"),
    ("Trạng thái hiện tại", "16/16 test đạt; API và Streamlit chạy; Docker chưa được kiểm chứng trên máy có Docker"),
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    font(paragraph.add_run(f"{label}: "), size=10.5, bold=True)
    font(paragraph.add_run(value), size=10.5)

rule = document.add_paragraph()
rule.paragraph_format.space_after = Pt(10)
ppr = rule._p.get_or_add_pPr()
border = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:color"), BLUE)
border.append(bottom)
ppr.append(border)

callout(
    document,
    "Kết luận điều hành",
    "Dự án giải quyết đúng hướng bài toán ước lượng giá đăng tham khảo, nhưng chưa đủ bằng chứng để gọi là production-ready. "
    "Ba ưu tiên bắt buộc là ẩn danh dữ liệu, sửa nhận diện tài sản trùng và tách validation khỏi calibration.",
    PALE_YELLOW,
)

document.add_heading("1. Bức tranh tổng thể và thứ tự ưu tiên", level=1)
add_table(
    document,
    ["Tầng", "Hiện trạng", "Mục tiêu", "Ưu tiên"],
    [
        ("Problem", "Đúng hướng nhưng phạm vi dữ liệu hẹp", "Định nghĩa rõ người dùng, đầu ra và giới hạn", "P1"),
        ("AI/ML correctness", "Có leakage risk; calibration dùng hai mục đích", "Split đúng, baseline rõ, coverage đáng tin", "P0"),
        ("Software Engineering", "Code tốt; Docker và integration test còn thiếu", "Clone/build/test nhất quán", "P1"),
        ("Production / Business", "Chưa load test; privacy đáng lo", "An toàn dữ liệu và benchmark được", "P0"),
    ],
    [1.25, 2.1, 2.35, 0.8],
)
document.add_paragraph("Quy ước ưu tiên: P0 = bắt buộc trước khi public; P1 = quan trọng; P2 = nâng cao sau khi demo ổn định.")

document.add_heading("2. Tầng Problem — xác nhận đúng bài toán", level=1)
document.add_heading("2.1 Định nghĩa lại bài toán", level=2)
document.add_paragraph(
    "Mục tiêu phù hợp nhất với dữ liệu hiện có là ước lượng giá đăng tham khảo cho nhà riêng và căn hộ trong phạm vi dữ liệu nghiên cứu tại TP.HCM. "
    "Không định vị mô hình như công cụ xác định giá giao dịch, thẩm định tài sản hoặc ra quyết định tín dụng."
)
add_table(
    document,
    ["Thành phần", "Định nghĩa nên dùng"],
    [
        ("Người dùng", "Người mua, người bán và nhà phân tích cần tham khảo nhanh"),
        ("Input", "Loại hình, khu vực, diện tích, phòng, tầng, kích thước, vị trí, tiện ích"),
        ("Output", "Giá điểm, khoảng dự báo, cảnh báo ngoài phạm vi, SHAP và trung vị phân khúc"),
        ("Đơn vị", "Triệu VND; phải thể hiện rõ khi chuyển sang tỷ VND trên UI"),
        ("Không sử dụng", "Thẩm định pháp lý, duyệt vay, giao dịch tự động hoặc cam kết lợi nhuận"),
    ],
    [1.35, 5.15],
)
document.add_heading("2.2 Tiêu chí nghiệm thu tầng Problem", level=2)
checklist(document, [
    "README ghi rõ đây là giá đăng tham khảo, không phải giá giao dịch.",
    "UI và API dùng cùng đơn vị, cùng tên trường và cùng disclaimer.",
    "Danh sách loại hình/khu vực trên UI được lấy từ model artifact, không phải danh mục rộng hơn dữ liệu train.",
    "Có ví dụ người dùng thực tế và trường hợp không nên dùng mô hình.",
    "Có mục Hạn chế riêng, dễ nhìn thấy trong README.",
])

document.add_heading("3. Tầng AI/ML correctness", level=1)
document.add_heading("3.1 Sửa nhận diện bất động sản trùng — P0", level=2)
callout(
    document,
    "Vấn đề",
    "property_group_id hiện ưu tiên Listing ID. Một căn nhà do nhiều môi giới đăng với ID khác nhau vẫn có thể rơi vào train và test. "
    "Khi mỗi ID gần như duy nhất, grouped split thực tế suy biến thành temporal row split.",
    PALE_RED,
)
document.add_paragraph("Cần tạo chữ ký tài sản độc lập với tin đăng. Không dùng Price, Listing ID, ngày đăng hoặc tên môi giới.")
code_block(document, """
def make_property_signature(df: pd.DataFrame) -> pd.Series:
    signature = pd.DataFrame({
        "location": df["Location"].map(_normalize_text),
        "property_type": df["Property Type"].map(_normalize_text),
        "area": pd.to_numeric(df["Area"], errors="coerce").round(0),
        "width": pd.to_numeric(df["Width"], errors="coerce").round(1),
        "length": pd.to_numeric(df["Length"], errors="coerce").round(1),
        "bedrooms": df["Bedrooms"],
        "bathrooms": df["Bathrooms"],
        "latitude": pd.to_numeric(df["Latitude"], errors="coerce").round(4),
        "longitude": pd.to_numeric(df["Longitude"], errors="coerce").round(4),
    })
    return pd.util.hash_pandas_object(
        signature.fillna("missing"), index=False
    ).astype(str)
""")
add_number(document, "Bước 1", "Chuẩn hóa Unicode, khoảng trắng, chữ thường và cách viết quận/huyện.")
add_number(document, "Bước 2", "Làm tròn số đo để các sai lệch nhập liệu nhỏ không tạo nhóm mới.")
add_number(document, "Bước 3", "Dùng tọa độ đã làm tròn khi có; dùng địa chỉ chuẩn hóa khi thiếu tọa độ.")
add_number(document, "Bước 4", "Kiểm tra phân phối kích thước nhóm trước khi chia tập.")
add_number(document, "Bước 5", "Chia toàn bộ nhóm, sau đó ánh xạ các dòng về đúng tập.")
checklist(document, [
    "Cùng địa chỉ nhưng Listing ID khác vẫn cùng property_group_id.",
    "Cùng tài sản nhưng thay đổi giá vẫn cùng nhóm.",
    "Hai căn khác nhau trong cùng quận không bị gộp sai.",
    "Có nhóm thực sự chứa nhiều hơn một tin đăng.",
    "Train/validation/calibration/test giao rỗng theo property_group_id.",
])

document.add_heading("3.2 Tách Validation và Calibration — P0", level=2)
document.add_paragraph(
    "Dự án đã tách Validation khỏi Calibration: Extra Trees được chọn trên Validation, còn Calibration chỉ dùng để tạo conformal quantile. "
    "Cấu trúc này cần được giữ nguyên trong các thử nghiệm tiếp theo."
)
add_table(
    document,
    ["Tập", "Tỷ lệ đề xuất", "Mục đích"],
    [
        ("Train", "60%", "Fit preprocessing và model"),
        ("Validation", "15%", "Chọn thuật toán, tham số và so sánh baseline"),
        ("Calibration", "10%", "Chỉ tính conformal quantile"),
        ("Test", "15%", "Đánh giá cuối cùng đúng một lần"),
    ],
    [1.35, 1.2, 3.95],
)
code_block(document, """
train_end = int(number_of_groups * 0.60)
validation_end = int(number_of_groups * 0.75)
calibration_end = int(number_of_groups * 0.85)
""")
callout(
    document,
    "Lưu ý cỡ mẫu",
    "707 dòng sau làm sạch là khá ít cho bốn tập. Không nên chạy tìm kiếm hyperparameter lớn. Báo cáo số nhóm và số dòng trong từng tập, "
    "đồng thời cảnh báo nếu validation, calibration hoặc test quá nhỏ.",
)

document.add_heading("3.3 Làm rõ confidence và conformal interval", level=2)
document.add_paragraph(
    "Các ngưỡng relative_width 0,4 và 0,8 là heuristic, không phải xác suất mô hình đúng. Nên đổi confidence thành reliability_level hoặc data_reliability."
)
add_table(
    document,
    ["Nhãn", "Điều kiện", "Cách diễn giải"],
    [
        ("High", "Dữ liệu đầy đủ, trong phạm vi, interval tương đối hẹp", "Chỉ báo vận hành thuận lợi"),
        ("Medium", "Thiếu một số trường hoặc interval trung bình", "Cần tham khảo thêm dữ liệu"),
        ("Low", "Ngoài phạm vi hoặc interval rộng", "Không nên dựa vào giá điểm"),
    ],
    [1.0, 2.6, 2.9],
)
callout(
    document,
    "Câu chữ đúng",
    "Khoảng dự báo được hiệu chỉnh với mục tiêu bao phủ 80%; high/medium/low là chỉ báo heuristic dựa trên chất lượng dữ liệu và độ rộng khoảng dự báo.",
    PALE_GREEN,
)

document.add_heading("3.4 Hoàn thiện metric và baseline", level=2)
add_bullet(document, "Giữ MAE, RMSE, R², MAPE và interval coverage.")
add_bullet(document, "Bổ sung Median Absolute Error, mean interval width, relative interval width và coverage gap.")
add_bullet(document, "Báo cáo MAE/coverage theo loại hình, quận và nhóm giá; chỉ báo nhóm có ít nhất 20 mẫu.")
add_bullet(document, "Dùng MAE validation làm tiêu chí chọn model và ghi rõ vì sao; không chọn bằng test.")
code_block(document, """
metrics["median_ae_million"] = float(median_absolute_error(actual, prediction))
metrics["mean_interval_width_million"] = float(np.mean(upper_bound - lower_bound))
metrics["coverage_gap"] = float(interval_coverage - target_coverage)

if extra_trees_validation_mae >= baseline_validation_mae:
    raise RuntimeError("Extra Trees không cải thiện MAE so với baseline.")
""")

document.add_heading("4. Dữ liệu và quyền riêng tư", level=1)
document.add_heading("4.1 Ẩn danh dữ liệu mẫu — P0", level=2)
callout(
    document,
    "Rủi ro",
    "CSV mẫu có thể chứa tên môi giới, số điện thoại, avatar, mô tả và địa chỉ chi tiết. Dù lấy từ tin công khai, đây vẫn là dữ liệu có khả năng nhận dạng cá nhân.",
    PALE_RED,
)
add_table(
    document,
    ["Nhóm dữ liệu", "Xử lý đề xuất"],
    [
        ("Agent Name, Avatar, Agent Role", "Xóa khỏi bản dữ liệu public"),
        ("Số điện thoại trong Description/Title", "Thay bằng [PHONE] hoặc xóa trường văn bản"),
        ("Địa chỉ chi tiết", "Chỉ giữ quận/khu vực đã chuẩn hóa"),
        ("Latitude/Longitude", "Làm tròn 3 chữ số hoặc chỉ giữ khoảng cách CBD"),
        ("Listing ID", "Hash hoặc tạo ID mới; không công bố ID nguồn nếu không cần"),
    ],
    [2.4, 4.1],
)
code_block(document, """
PHONE_PATTERN = r"(?:\\+?84|0)\\d{8,10}"
df["Description"] = (
    df["Description"].fillna("")
    .str.replace(PHONE_PATTERN, "[PHONE]", regex=True)
)
df["Latitude"] = df["Latitude"].round(3)
df["Longitude"] = df["Longitude"].round(3)
""")
document.add_paragraph("Kiểm tra trước khi commit:")
code_block(document, """
rg -n "\\b0[0-9]{8,10}\\b" data
rg -n "Agent Name|Avatar|Phone" data
""")

document.add_heading("4.2 Data card có khả năng truy vết", level=2)
document.add_paragraph("Dự án hiện ghi rõ đường đi 2.500 dòng thô → 707 dòng sau làm sạch bằng số lượng loại theo từng lý do.")
add_bullet(document, "Số dòng bị loại vì loại hình không hỗ trợ.")
add_bullet(document, "Số dòng bị loại vì giá/diện tích/đơn giá bất hợp lý.")
add_bullet(document, "Tỷ lệ thiếu theo từng cột và tỷ lệ tọa độ hợp lệ.")
add_bullet(document, "Số nhóm trùng, kích thước nhóm lớn nhất và duplicate group percent.")
add_bullet(document, "Phân vị target, diện tích và số dòng theo loại hình/quận.")

document.add_heading("5. Tầng Software Engineering", level=1)
document.add_heading("5.1 Bổ sung integration và edge-case tests", level=2)
add_table(
    document,
    ["Nhóm test", "Trường hợp cần thêm"],
    [
        ("Artifact", "Thiếu model, model sai cấu trúc, version không tương thích"),
        ("Dữ liệu", "Rỗng, dưới ngưỡng, toàn NaN, category chưa từng thấy"),
        ("Prediction", "Cận không âm, upper ≥ lower, SHAP tối đa 5 phần tử"),
        ("API", "POST /predict dùng model thật, không mock; schema và lỗi 503"),
        ("Split", "Bốn tập giao rỗng và thứ tự thời gian đúng"),
        ("Conformal", "Residual rỗng, coverage ngoài (0,1), coverage gap được ghi"),
    ],
    [1.35, 5.15],
)
code_block(document, """
def test_real_model_prediction_schema():
    result = predict_one({
        "Property Type": "Nhà riêng",
        "location_area": "Quận 1",
        "Area": 80,
        "Bedrooms": 3,
    })
    assert result["predicted_price_million"] >= 0
    assert result["upper_bound_million"] >= result["lower_bound_million"]
    assert len(result["top_contributions"]) <= 5
""")

document.add_heading("5.2 Sửa Docker healthcheck — P1", level=2)
document.add_paragraph("Image python:3.11-slim không mặc định có curl, trong khi Compose dùng curl để kiểm tra sức khỏe.")
code_block(document, """
healthcheck:
  test:
    [
      "CMD",
      "python",
      "-c",
      "import urllib.request; urllib.request.urlopen('http://localhost:8000/health')"
    ]
  interval: 30s
  timeout: 10s
  retries: 3
  start_period: 20s
""")

document.add_heading("5.3 Không huấn luyện trong Docker build", level=2)
document.add_paragraph(
    "RUN python -m src.train làm build chậm và khiến model phụ thuộc quá trình build. Với portfolio, nên commit model artifact đã kiểm chứng và chỉ validate artifact khi build."
)
code_block(document, """
# Bỏ khỏi Dockerfile:
# RUN python -m src.train

# Thay bằng kiểm tra artifact:
RUN python -c "from src.predict import load_model; load_model()"
""")

document.add_heading("5.4 Chạy container bằng non-root", level=2)
code_block(document, """
RUN useradd --create-home appuser \\
    && chown -R appuser:appuser /app
USER appuser
""")

document.add_heading("5.5 Checklist chạy từ clone sạch", level=2)
code_block(document, """
git clone https://github.com/haminhthong/hcmc-real-estate-price-intelligence.git
cd hcmc-real-estate-price-intelligence
python -m venv .venv
pip install -r requirements.txt
python -m src.train
python -m src.evaluate
pytest -q
uvicorn api.main:app
streamlit run app/streamlit_app.py
""")
checklist(document, [
    "Không cần mở notebook hoặc sửa đường dẫn.",
    "Không tạo NaN/Infinity sau transform.",
    "API /health, /model-info và /predict trả đúng schema.",
    "Docker build và Compose healthcheck đạt trên máy có Docker.",
    "CI chạy trên checkout sạch và không dựa vào cache máy cá nhân.",
])

document.add_heading("6. Tầng Production / Business value", level=1)
document.add_heading("6.1 Tách SHAP khỏi dự báo mặc định", level=2)
document.add_paragraph("SHAP được tính cho mọi request, gây tốn CPU và làm giảm throughput. Có hai lựa chọn:")
add_bullet(document, "Thêm include_explanation=false vào /predict và chỉ tính SHAP khi được yêu cầu.")
add_bullet(document, "Tốt hơn: tách POST /predict và POST /explain, áp rate limit thấp hơn cho /explain.")
code_block(document, """
contributions = (
    explain_top_features(model_package, feature_frame)
    if include_explanation
    else []
)
""")

document.add_heading("6.2 Chuẩn bị cho 100 người dùng", level=2)
add_number(document, "Bước 1", "Chạy 2 Uvicorn workers và theo dõi RAM vì mỗi worker tải một bản model.")
add_number(document, "Bước 2", "Load test bằng Locust hoặc k6 với ramp-up, không gửi 100 request cùng lúc ngay lập tức.")
add_number(document, "Bước 3", "Đo p50, p95, p99 latency; throughput; error rate; CPU và RAM.")
add_number(document, "Bước 4", "Tách luồng có SHAP và không SHAP để biết chi phí giải thích.")
add_number(document, "Bước 5", "Chỉ tuyên bố hỗ trợ 100 users sau khi có báo cáo benchmark lặp lại được.")
add_table(
    document,
    ["Kịch bản", "Tỷ lệ", "Mục tiêu ban đầu"],
    [
        ("/predict không SHAP", "80%", "p95 < 1 giây"),
        ("/predict hoặc /explain có SHAP", "15%", "p95 < 3 giây"),
        ("/model-info", "5%", "p95 < 300 ms"),
        ("Toàn hệ thống", "100 users", "Error rate < 1%, không OOM"),
    ],
    [2.35, 1.0, 3.15],
)

document.add_heading("6.3 Security tối thiểu cho demo public", level=2)
add_bullet(document, "Rate limit /predict khoảng 30 request/phút/IP và /explain khoảng 5 request/phút/IP.")
add_bullet(document, "Giới hạn kích thước request, timeout và số worker phù hợp RAM.")
add_bullet(document, "Chỉ bật CORS cho domain Streamlit cụ thể nếu trình duyệt thực sự gọi API.")
add_bullet(document, "Không tải file joblib từ nguồn không tin cậy vì joblib/pickle có thể thực thi mã khi deserialize.")
add_bullet(document, "Không ghi địa chỉ đầy đủ, tọa độ chính xác hoặc nội dung mô tả vào log.")

document.add_heading("6.4 Logging và monitoring", level=2)
add_table(
    document,
    ["Nên ghi", "Không nên ghi"],
    [
        ("request_id, model_version", "Tên, số điện thoại"),
        ("status_code, response_time_ms", "Địa chỉ đầy đủ"),
        ("warnings_count, explanation_enabled", "Tọa độ chính xác"),
        ("out_of_distribution_count", "Description thô"),
    ],
    [3.25, 3.25],
)

document.add_heading("7. README và giá trị portfolio", level=1)
document.add_heading("7.1 Giảm nội dung mang tính trình diễn", level=2)
add_bullet(document, "Giảm các câu production-ready, 1-click và triệt tiêu hoàn toàn leakage/skew nếu chưa có bằng chứng.")
add_bullet(document, "Chuyển Điểm sáng CV và Góc phỏng vấn xuống cuối hoặc sang tài liệu riêng.")
add_bullet(document, "Không viết đảm bảo coverage 80% khi test coverage hiện chỉ là 71,03%.")
add_bullet(document, "Dùng số liệu kiểm chứng thay cho tính từ như chính xác, mạnh mẽ hoặc production-grade.")
callout(
    document,
    "Câu nên dùng",
    "Extra Trees giảm test MAE khoảng 20,5% so với median baseline; tuy nhiên R² chỉ đạt 0,232 và MAPE khoảng 53,74%, "
    "vì vậy kết quả chỉ phù hợp cho tham khảo trong phạm vi dữ liệu nghiên cứu.",
    PALE_GREEN,
)

document.add_heading("7.2 Mục Hạn chế nên có", level=2)
checklist(document, [
    "Dữ liệu là giá đăng, không phải giá giao dịch.",
    "Chỉ còn 707 mẫu sau làm sạch.",
    "R² 0,232; MAPE 53,74%; coverage 71,03% so với mục tiêu 80%.",
    "Nhiều bản ghi thiếu tọa độ và địa chỉ chưa đủ chi tiết.",
    "Nhận diện cùng một tài sản giữa nhiều môi giới chưa hoàn hảo.",
    "Kết quả không thay thế thẩm định chuyên nghiệp.",
])

document.add_heading("8. Lộ trình triển khai đề xuất", level=1)
add_table(
    document,
    ["Giai đoạn", "Công việc", "Đầu ra nghiệm thu"],
    [
        ("1 — Public-safe", "Ẩn danh data; limitation; kiểm tra secret/PII", "Không còn dữ liệu nhận dạng"),
        ("2 — ML-correct", "Group signature; 4-way split; metric/quality gate", "Leakage tests và coverage report"),
        ("3 — Reproducible", "Integration tests; Docker; clone sạch", "CI + Docker health đều đạt"),
        ("4 — Scalable demo", "Tách SHAP; workers; rate limit; load test", "Báo cáo 100 users"),
        ("5 — Advanced", "Monitoring; MLflow nếu cần", "Có trace và lịch sử model"),
    ],
    [1.25, 3.25, 2.0],
)

document.add_heading("9. Definition of Done", level=1)
checklist(document, [
    "Không có dữ liệu cá nhân trong data/sample.",
    "Các property_group_id không trùng giữa train/validation/calibration/test.",
    "Calibration không được dùng để chọn model.",
    "Model vượt baseline theo metric đã công bố.",
    "Coverage, coverage gap và interval width được báo cáo.",
    "Ruff, pytest, train và evaluate đều đạt trên checkout sạch.",
    "API và Streamlit smoke test đạt.",
    "Docker build, healthcheck và Compose đều đạt.",
    "README phản ánh đúng metrics, phạm vi và limitations.",
    "Có load-test report trước khi tuyên bố hỗ trợ nhiều người dùng.",
])

document.add_heading("10. Kết luận", level=1)
document.add_paragraph(
    "Không cần bổ sung Kubernetes, Airflow hoặc microservice. Giá trị portfolio sẽ tăng mạnh hơn khi dự án chứng minh được ba điều: "
    "dữ liệu an toàn, đánh giá ML không rò rỉ và hệ thống có thể chạy lặp lại từ checkout sạch. Sau khi các yêu cầu P0/P1 đạt, "
    "FastAPI + Streamlit + Docker + CI đã đủ để thể hiện năng lực Data, AI và Software Engineering."
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
document.save(OUTPUT)
